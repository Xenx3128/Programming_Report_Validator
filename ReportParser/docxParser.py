import os
from pathlib import Path

import docx.document
import xml.etree.ElementTree as ET
from docx.api import Document
from docx.document import Document as doctwo
from docx.shared import Pt, Cm, Emu
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_UNDERLINE, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT as WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENTATION
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.comments import Comment
from io import StringIO
import re

from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Length
from docx.text.run import Run

from default_settings import *
from constants import *
import inspect

import numbers


class PCommentGeneral:
    def __init__(self, topic=None, correct=None, received=None):
        self.topic = topic
        self.correct = correct
        self.received = received


class PCommentSpecial:
    def __init__(self, topic=None, text=None):
        self.topic = topic
        self.text = text
        
    
class BaseChecklist:
    def set_settings(self, settings: dict):
        for par, value in settings.items():
            try:
                setattr(self, par, value)
            except:
                continue


class ParagraphChecklist(BaseChecklist):
    def __init__(self):
        super().__init__()
        self.format_regex = None
        self.font_name = None
        self.font_size = None
        self.font_bald = None
        self.font_italic = None
        self.font_underline = None
        self.font_color = None  # !!!
        self.font_highlight = None  # !!!
        self.alignment = None
        self.keep_with_next = None
        self.page_break_before = None
        self.space_before = None
        self.space_after = None
        self.left_indent = None
        self.right_indent = None
        self.first_line_indent = None
        self.line_spacing = None
        self.header_numbering = None



class ListChecklist(ParagraphChecklist):
    def __init__(self):
        super().__init__()
        self.left_indent_base = None
        self.left_indent_mod = None



class TableParagraphChecklist(ParagraphChecklist):
    def __init__(self):
        super().__init__()
        self.vert_alignment = None


class ImageChecklist(BaseChecklist):
    def __init__(self):
        super().__init__()
        self.alignment = None
        self.keep_with_next = None
        self.space_before = None
        self.space_after = None
        self.first_line_indent = None
        self.line_spacing = None


class MarginsChecklist(BaseChecklist):
    def __init__(self):
        super().__init__()
        self.top_margin = None
        self.bottom_margin = None
        self.left_margin = None
        self.right_margin = None
        self.orientation = None


class DocumentParser:
    def __init__(self):
        self.text_checklist = ParagraphChecklist()
        self.heading1_checklist = ParagraphChecklist()
        self.heading2_checklist = ParagraphChecklist()
        self.heading3_checklist = ParagraphChecklist()
        self.bullet_list_checklist = ListChecklist()
        self.nubered_list_checklist = ListChecklist()
        self.table_headings_checklist = TableParagraphChecklist()
        self.table_text_checklist = TableParagraphChecklist()
        self.table_name_checklist = ParagraphChecklist()
        self.text_after_table_checklist = ParagraphChecklist()
        self.image_checklist = ImageChecklist()
        self.image_name_checklist = ParagraphChecklist()
        self.text_before_list_checklist = ParagraphChecklist()
        self.margins_checklist = MarginsChecklist()
        self.set_settings(default_text_checklist, default_heading1_checklist, default_heading2_checklist,
                          default_heading3_checklist, default_table_name_checklist, default_table_headings_checklist,
                          default_table_text_checklist, default_bullet_list_checklist, default_numbered_list_checklist, default_margins_checklist,
                          default_image_checklist, default_image_name_checklist, default_text_before_list_checklist, default_text_after_table_checklist)

        self.enable_optional_settings = {
            "table_headings_top": True,
            "table_headings_left": False,
            "table_title": True,
            "paragraph_after_table": True,
            "enable_pic_title": True
        }

        
        self.doc_comments = {}

    def set_enable_optional_settings(self, elements: dict):
        if elements is not None:
            for par, value in elements.items():
                if par in self.enable_optional_settings:
                    self.enable_optional_settings[par] = value

    def set_settings(self, text_check=None, h1_check=None, h2_check=None, h3_check=None, table_name_check=None,
                     table_heading_check=None, table_text_check=None, bullet_list_check=None, numbered_list_check=None,
                     page_check=None, image_check=None, image_name_check=None, text_before_list_check=None, text_after_table_check=None, enable_optional_settings=None):
        if isinstance(text_check, dict):
            self.text_checklist.set_settings(text_check)
        if isinstance(h1_check, dict):
            self.heading1_checklist.set_settings(h1_check)
        if isinstance(h2_check, dict):
            self.heading2_checklist.set_settings(h2_check)
        if isinstance(h3_check, dict):
            self.heading3_checklist.set_settings(h3_check)
        if isinstance(table_name_check, dict):
            self.table_name_checklist.set_settings(table_name_check)
        if isinstance(table_heading_check, dict):
            self.table_headings_checklist.set_settings(table_heading_check)
        if isinstance(table_text_check, dict):
            self.table_text_checklist.set_settings(table_text_check)
        if isinstance(bullet_list_check, dict):
            self.bullet_list_checklist.set_settings(bullet_list_check)
        if isinstance(numbered_list_check, dict):
            self.nubered_list_checklist.set_settings(numbered_list_check)
        if isinstance(page_check, dict):
            self.margins_checklist.set_settings(page_check)
        if isinstance(image_check, dict):
            self.image_checklist.set_settings(image_check)
        if isinstance(image_name_check, dict):
            self.image_name_checklist.set_settings(image_name_check)
        if isinstance(text_before_list_check, dict):
            self.text_before_list_checklist.set_settings(text_before_list_check)
        if isinstance(text_after_table_check, dict):
            self.text_after_table_checklist.set_settings(text_after_table_check)
        if isinstance(enable_optional_settings, dict):
            for par, value in enable_optional_settings.items():
                if par in self.enable_optional_settings:
                    self.enable_optional_settings[par] = value

    
    @staticmethod
    def get_error_comment(checklist, received: dict, epsilon=0.1):
        attributes = inspect.getmembers(checklist, lambda a: not (inspect.isroutine(a)))
        expected = {a[0]: a[1] for a in attributes if not (a[0].startswith('__') and a[0].endswith('__'))}
        comments = []
        
        for key, val in expected.items():
            comparison_res = True
            comment = []
            if key in received.keys() and val is not None:
                # Отдельная обработка формата через RegEx
                if key == 'format_regex':
                    comparison_res = re.fullmatch(fr'{REGEX_TRANSFORM[val]}', received[key])
                    if not comparison_res:

                        comments.append(PCommentGeneral(PARAM_TO_COMMENT[key], val, received[key]))
                        continue
                if "left_indent_base" in expected and "is_list" in received and received["list_level"] > 0 and key == "left_indent":
                    expected_indent = expected["left_indent_base"] + expected["left_indent_mod"] * (received["list_level"] - 1)
                    comparison_res = abs(received["left_indent"] - expected_indent) <= epsilon
                    if not comparison_res:
                        comments.append(PCommentGeneral(PARAM_TO_COMMENT[key], expected_indent, received[key]))
                        continue
                
                if isinstance(val, bool):
                    if received[key] == True:
                        comparison_res = val == True
                    else:
                        comparison_res = val in (False, None)
                elif isinstance(val, numbers.Number) and isinstance(received[key], numbers.Number):
                    comparison_res = abs(val - received[key]) <= epsilon
                else:
                    comparison_res = val == received[key]
                if not comparison_res:
                    if key in SPECIAL_COMMENTS:
                        comment = PCommentSpecial(topic=PARAM_TO_COMMENT[key], text=SPECIAL_COMMENTS[key])
                    else:
                        comment = PCommentGeneral(topic=PARAM_TO_COMMENT[key])
                        # У выравниваний значения приравниваются к 0..3 через константы
                        if key == "alignment":
                            comment.correct = ALIGNMENT_TO_COMMENT[val]
                            comment.received = ALIGNMENT_TO_COMMENT[received[key]]
                        elif key == "vert_alignment":
                            comment.correct = VERT_ALIGNMENT_TO_COMMENT[val]
                            comment.received = VERT_ALIGNMENT_TO_COMMENT[received[key]] 
                        elif key == "orientation":
                            comment.correct = ORIENTATION_TO_COMMENT[val]
                            comment.received = ORIENTATION_TO_COMMENT[received[key]]
                        else:
                            if val is None or isinstance(val, bool):
                                comment.correct = VAR_TO_COMMENT[val] 
                            else:
                                comment.correct = val
                            if received is None or isinstance(received[key], bool):
                                comment.received = VAR_TO_COMMENT[received[key]]
                            else:
                                comment.received = received[key]
                    comments.append(comment)
        return comments
    
    @staticmethod
    def create_comment(document, runs, element, comments: list):
        if len(comments) > 0 and len(runs) > 0:
            comment = document.add_comment(
                runs=runs,
                text='',
                author=element
            )
            if element == 'Таблица':
                cmt_para = comment.add_paragraph()
                cmt_para.add_run("Выведена только первая ошибка в каждой категории\n").bold = True 
            for cmt in comments:
                if isinstance(cmt, PCommentGeneral):
                    cmt_para = comment.add_paragraph()
                    cmt_para.add_run(f"{cmt.topic}: ").bold = True 
                    cmt_para.add_run(f"{cmt.received} | ")
                    cmt_para.add_run(f"({cmt.correct}).")
                elif isinstance(cmt, PCommentSpecial):
                    cmt_para = comment.add_paragraph()
                    cmt_para.add_run(f"{cmt.topic}: ").bold = True 
                    cmt_para.add_run(f"{cmt.text}")
            return comment

 
    def get_paragraph_properties(self, document, paragraph):

        if paragraph is None:
            return None
                
        def twips_to_emu(twips):
            return int(twips * 914400 / 1440)
        
        
        def get_numbering_properties(paragraph, document):

            props = {}

            def extract_numPr_from_pPr(pPr_element):
                """Общий экстрактор numPr → ilvl, num_id"""
                if pPr_element is None:
                    return None, None
                numPr = pPr_element.find(qn('w:numPr'))
                if numPr is None:
                    return None, None
                ilvl_el = numPr.find(qn('w:ilvl'))
                numId_el = numPr.find(qn('w:numId'))
                ilvl = int(ilvl_el.get(qn('w:val'), '0')) if ilvl_el is not None else 0
                num_id = int(numId_el.get(qn('w:val'))) if numId_el is not None else None
                return ilvl, num_id

            # Прямое форматирование параграфа
            pPr_direct = paragraph._element.find(qn('w:pPr'))
            ilvl, num_id = extract_numPr_from_pPr(pPr_direct)
            if num_id:
                pass
            else:
                # Цепочка стилей
                style = paragraph.style
                visited = set()
                while style and style.name not in visited:
                    visited.add(style.name)
                    try:
                        style_pPr = style._element.find(qn('w:pPr'))
                        ilvl, num_id = extract_numPr_from_pPr(style_pPr)
                        if num_id is not None:
                            break
                    except AttributeError:
                        pass
                    style = style.base_style

            if not num_id:
                return props

            try:
                numbering_part = document.part.numbering_part
                if not numbering_part or not numbering_part._element:
                    return props

                num_el = numbering_part._element

                # lvlOverride (приоритетнее abstract)
                lvl_nodes = num_el.xpath(
                    f'.//w:num[@w:numId="{num_id}"]/w:lvlOverride[@w:ilvl="{ilvl}"]/w:lvl'
                )

                if not lvl_nodes:
                    # abstractNum
                    abstract_refs = num_el.xpath(
                        f'.//w:num[@w:numId="{num_id}"]/w:abstractNumId'
                    )
                    if not abstract_refs:
                        return props
                    abstract_id = abstract_refs[0].get(qn('w:val'))
                    if not abstract_id:
                        return props

                    lvl_nodes = num_el.xpath(
                        f'.//w:abstractNum[@w:abstractNumId="{abstract_id}"]/w:lvl[@w:ilvl="{ilvl}"]'
                    )

                if not lvl_nodes:
                    return props

                lvl = lvl_nodes[0]

                numFmt_el = lvl.find(qn('w:numFmt'))
                if numFmt_el is not None:
                    val = numFmt_el.get(qn('w:val'))
                    if val:
                        props['numbering_format'] = val
                        props['list_type'] = 'bulleted' if val == 'bullet' else 'numbered'

                lvl_pPr = lvl.find(qn('w:pPr'))
                if lvl_pPr is None:
                    return props

                # indents
                ind = lvl_pPr.find(qn('w:ind'))
                if ind is not None:
                    for side in ['left', 'start', 'right', 'end']:
                        val = ind.get(qn(f'w:{side}'))
                        if val is not None:
                            props[f'{side}_indent'] = Length(twips_to_emu(int(val)))

                    hanging = ind.get(qn('w:hanging'))
                    first_line = ind.get(qn('w:firstLine'))
                    if hanging is not None:
                        props['first_line_indent'] = Length(twips_to_emu(-int(hanging)))
                    elif first_line is not None:
                        props['first_line_indent'] = Length(twips_to_emu(int(first_line)))

                # alignment
                jc = lvl_pPr.find(qn('w:jc'))
                if jc is not None:
                    val = jc.get(qn('w:val'))
                    if val:
                        try:
                            props['alignment'] = getattr(WD_PARAGRAPH_ALIGNMENT, val.upper())
                        except AttributeError:
                            pass

                # spacing
                spacing = lvl_pPr.find(qn('w:spacing'))
                if spacing is not None:
                    for attr in ['before', 'after']:
                        val = spacing.get(qn(f'w:{attr}'))
                        if val is not None:
                            props[f'space_{attr}'] = Length(twips_to_emu(int(val)))
                    line = spacing.get(qn('w:line'))
                    line_rule = spacing.get(qn('w:lineRule'))
                    if line is not None:
                        props['line_spacing_raw'] = (int(line), line_rule)

                # flags
                for flag_name, prop_key in [('keepNext', 'keep_with_next'), ('pageBreakBefore', 'page_break_before')]:
                    flag_el = lvl_pPr.find(qn(f'w:{flag_name}'))
                    if flag_el is not None:
                        val = flag_el.get(qn('w:val'))
                        props[prop_key] = val in (None, '1', 'true', 'on')

            except Exception:
                pass

            return props

        def resolve_para_prop(prop_name, default=None):

            # Прямое
            try:
                fmt = paragraph.paragraph_format
                value = getattr(fmt, prop_name, None)
                if value is not None:
                    return value
            except Exception:
                pass

            # Numbering
            numbering_props = get_numbering_properties(paragraph, document)
            if prop_name in numbering_props:
                return numbering_props[prop_name]

            if prop_name == 'line_spacing' and 'line_spacing_raw' in numbering_props:
                raw, rule = numbering_props['line_spacing_raw']
                if rule in (None, 'auto'):
                    return raw / 240.0  # multiplier
                elif rule in ('exactly', 'atLeast'):
                    return Length(raw).pt
                return raw / 240.0

            # Свойства стиля
            style = paragraph.style
            visited = set()
            while style and style.name not in visited:
                visited.add(style.name)
                try:
                    value = getattr(style.paragraph_format, prop_name, None)
                    if value is not None:
                        return value
                except Exception:
                    pass
                style = style.base_style

            return default

        
        is_list = self.is_list_item(paragraph)
        
        if is_list[0]:
            pass
        
        numbering_props = get_numbering_properties(paragraph, document)
        list_type = numbering_props.get('list_type')
        numbering_format = numbering_props.get('numbering_format')

        alignment = resolve_para_prop("alignment", WD_PARAGRAPH_ALIGNMENT.LEFT)
        keep_with_next = bool(resolve_para_prop("keep_with_next"))
        page_break_before = bool(resolve_para_prop("page_break_before"))

        left_l   = resolve_para_prop("left_indent")
        right_l  = resolve_para_prop("right_indent")
        first_l  = resolve_para_prop("first_line_indent")

        left_indent     = round(left_l.cm, 2)   if left_l  is not None else 0.0
        right_indent    = round(right_l.cm, 2)  if right_l is not None else 0.0
        first_line_indent = round(first_l.cm, 2) if first_l is not None else 0.0
        
        if first_line_indent < 0:  # выставлен выступ, а не отступ
            left_indent += first_line_indent

        before_l = resolve_para_prop("space_before")
        after_l  = resolve_para_prop("space_after")

        space_before = round(before_l.pt, 1) if before_l is not None else 0.0
        space_after  = round(after_l.pt, 1)  if after_l  is not None else 0.0

        line_raw = resolve_para_prop("line_spacing")
        
        if line_raw is None:
            line_spacing = 1.0
        elif hasattr(line_raw, 'pt'):
            line_spacing = round(line_raw.pt / 12, 2)
        else:
            line_spacing = round(float(line_raw), 2)


        return {
            "alignment": alignment,
            "keep_with_next": keep_with_next,
            "page_break_before": page_break_before,

            "left_indent": left_indent,
            "right_indent": right_indent,
            "first_line_indent": first_line_indent,

            "space_before": space_before,
            "space_after": space_after,
            "line_spacing": line_spacing,

            "is_list": is_list[0],
            'list_level': is_list[1],
            
            "list_type": list_type,
            "numbering_format": numbering_format
        }
 
    def get_run_properties(self, document, paragraph, run):

        if run is None or not hasattr(run, '_element'):
            return None

        if not hasattr(self, '_docx_style_cache'):
            self._docx_style_cache = {}
            self._theme_cache = None
            self._doc_defaults = None


        def get_theme_font(font_type='minor'):
            if self._theme_cache is None:
                self._theme_cache = {'minor': {}, 'major': {}}
                try:
                    theme = document.part.theme_part.theme
                    fs = theme.themeElements.fontScheme
                    fgroup = getattr(fs, f'{font_type}Font', None)
                    if fgroup:
                        self._theme_cache[font_type] = {
                            'latin': getattr(fgroup.latin, 'typeface', None),
                            'hAnsi': getattr(getattr(fgroup, 'hAnsi', None), 'typeface', None),
                            'eastAsia': getattr(getattr(fgroup, 'eastAsia', None), 'typeface', None),
                            'cs': getattr(getattr(fgroup, 'cs', None), 'typeface', None),
                        }
                except Exception:
                    pass
                if not self._theme_cache[font_type]:
                    try:
                        el = document.part.theme_part.element
                        typ = 'minorFont' if font_type == 'minor' else 'majorFont'
                        res = el.xpath(f'.//a:fontScheme/a:{typ}/a:latin')
                        if res:
                            typeface = res[0].get('typeface')
                            self._theme_cache[font_type] = {'latin': typeface, 'hAnsi': typeface}
                    except Exception:
                        pass
                default = "Calibri" if font_type == 'minor' else "Calibri Light"
                if not self._theme_cache[font_type]:
                    self._theme_cache[font_type] = {'latin': default, 'hAnsi': default}
            return self._theme_cache[font_type]

        def resolve_theme_name(name):
            if not name:
                return "Calibri"
            name = str(name).strip()
            if not name:
                return "Calibri"
            lower = name.lower()
            minor_placeholders = {"+body", "body", "+mnlt", "mnlt", "minor", "minorhansi", "minorascii",
                                  "minorbidi", "minoreastasia", "asciitheme", "hansitheme", "cstheme"}
            major_placeholders = {"+heading", "heading", "+mnhansi", "mnhansi", "major", "majorhansi",
                                  "majorascii", "majorbidi", "majoreastasia"}
            is_major = lower in major_placeholders
            theme_group = get_theme_font('major' if is_major else 'minor')
            for key in ('hAnsi', 'ascii', 'latin', 'eastAsia', 'cs'):
                val = theme_group.get(key)
                if val and str(val).strip():
                    return str(val).strip()
            return "Calibri Light" if is_major else "Calibri"

        def get_color_value(color_format):
            if not color_format:
                return None
            try:
                if hasattr(color_format, 'rgb') and color_format.rgb and str(color_format.rgb) != '000000':
                    rgb = color_format.rgb
                    return f"#{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
                if hasattr(color_format, 'theme_color') and color_format.theme_color:
                    tc = color_format.theme_color
                    name = getattr(tc, 'name', str(tc)).lower()
                    tint = getattr(color_format, 'brightness', None)
                    return f"theme:{name}" + (f"(tint={tint:.2f})" if tint else "")
            except Exception:
                pass
            return None

        def get_prop_from_font(font_obj, prop_name):
            if not font_obj:
                return None
            try:
                if prop_name == "name":
                    raw = getattr(font_obj, 'name', None)
                    return raw if raw else None
                if prop_name == "size":
                    size = getattr(font_obj, 'size', None)
                    return size.pt if size else None
                if prop_name in ("bold", "italic", "strike", "subscript", "superscript",
                                 "small_caps", "all_caps"):
                    return getattr(font_obj, prop_name, None)
                if prop_name == "underline":
                    return getattr(font_obj, 'underline', None)
                if prop_name == "color":
                    return get_color_value(font_obj.color)
                if prop_name == "highlight":
                    return getattr(font_obj, 'highlight_color', None)
            except Exception:
                pass
            return None

        def get_style_chain_prop(style, prop_name):
            if not style:
                return None
            visited = set()
            current = style
            while current and current.name not in visited:
                visited.add(current.name)
                value = get_prop_from_font(current.font, prop_name)
                if value is not None and (prop_name != "name" or value):
                    return value
                current = getattr(current, 'base_style', None)
            return None

        def resolve_run_prop(prop_name, default=None):
            value = get_prop_from_font(run.font, prop_name)
            if value is not None and (prop_name != "name" or value):
                return value
            if run.style:
                value = get_style_chain_prop(run.style, prop_name)
                if value is not None and (prop_name != "name" or value):
                    return value
            try:
                linked = getattr(paragraph.style, 'linked_style', None)
                if linked:
                    value = get_style_chain_prop(linked, prop_name)
                    if value is not None and (prop_name != "name" or value):
                        return value
            except Exception:
                pass
            if paragraph.style:
                value = get_style_chain_prop(paragraph.style, prop_name)
                if value is not None and (prop_name != "name" or value):
                    return value

            if self._doc_defaults is None:
                try:
                    self._doc_defaults = document.styles.element.xpath(
                        './/w:docDefaults/w:rPrDefault/w:rPr'
                    )
                    self._doc_defaults = self._doc_defaults[0] if self._doc_defaults else None
                except Exception:
                    self._doc_defaults = None
            if self._doc_defaults is not None and prop_name == "name":
                try:
                    rFonts = self._doc_defaults.find(qn('w:rFonts'))
                    if rFonts is not None:
                        for attr in ['hAnsi', 'ascii', 'eastAsia', 'cs']:
                            val = rFonts.get(qn(f'w:{attr}'))
                            if val:
                                return val
                except Exception:
                    pass
            if prop_name == "name":
                return resolve_theme_name("+Body") or default
            return default
        
        
        font_name = resolve_run_prop("name", "Calibri")
        font_size = round(resolve_run_prop("size") or 11.0, 1)
        font_bold = bool(resolve_run_prop("bold"))
        font_italic = bool(resolve_run_prop("italic"))
        font_underline = resolve_run_prop("underline") not in (None, False, WD_UNDERLINE.NONE)
        font_color = resolve_run_prop("color")
        font_highlight = resolve_run_prop("highlight")
        font_strike = bool(resolve_run_prop("strike"))
        font_subscript = bool(resolve_run_prop("subscript"))
        font_superscript = bool(resolve_run_prop("superscript"))
        font_small_caps = bool(resolve_run_prop("small_caps"))
        font_all_caps = bool(resolve_run_prop("all_caps"))

        result = {
            "font_name": font_name,
            "font_size": font_size,
            "font_bold": font_bold,
            "font_italic": font_italic,
            "font_underline": font_underline,
            "font_color": font_color,
            "font_highlight": font_highlight,
            "font_strike": font_strike,
            "font_subscript": font_subscript,
            "font_superscript": font_superscript,
            "font_small_caps": font_small_caps,
            "font_all_caps": font_all_caps,
        }
            
        return result

    # This function extracts the tables and paragraphs from the document object
    @staticmethod
    def iter_block_items(parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph. *parent*
        would most commonly be a reference to a main Document object, but
        also works for a _Cell object, which itself can contain paragraphs and tables.
        """
        if isinstance(parent, doctwo):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    @staticmethod
    def set_comment_name(name, prev_name, is_locked):
        if not is_locked:
            return name
        return prev_name

    def parse_margins(self, document):
        sections = document.sections
        paragraph_stats = {}
        stats_to_compare = self.margins_checklist
        margins_comments = {}
        for section_cnt, section in enumerate(sections):
            paragraph_stats["top_margin"] = round(section.top_margin.cm, 2)
            paragraph_stats["bottom_margin"] = round(section.bottom_margin.cm, 2)
            paragraph_stats["left_margin"] = round(section.left_margin.cm, 2)
            paragraph_stats["right_margin"] = round(section.right_margin.cm, 2)
            paragraph_stats["orientation"] = section.orientation
            section_comments = self.get_error_comment(stats_to_compare, paragraph_stats)
            if len(section_comments) > 0:
                margins_comments[f"Раздел {section_cnt + 1}:"] = section_comments
        return margins_comments     


    @staticmethod
    def get_heading_level(paragraph):
        def extract_level_from_style(style):
            if not style:
                return None
            name = style.name or ""
            if name.startswith("Heading "):
                try:
                    return int(name.split()[-1])
                except (ValueError, IndexError):
                    pass
            style_id = getattr(style, 'style_id', None) or ""
            if style_id.startswith("Heading"):
                try:
                    return int(style_id.replace("Heading", "").strip())
                except ValueError:
                    pass
            if name.startswith("Heading") and len(name) > 7:
                try:
                    return int(name[7:])
                except ValueError:
                    pass
            return None

        level = extract_level_from_style(paragraph.style)
        if level is not None:
            return level

        current = paragraph.style
        depth = 0
        while current and depth < 20:
            level = extract_level_from_style(current)
            if level is not None:
                return level
            current = current.base_style
            depth += 1


        outline = paragraph._p.xpath('./w:pPr/w:outlineLvl/@w:val')
        if outline:
            try:
                lvl = int(outline[0]) + 1
                return lvl if 1 <= lvl <= 9 else None
            except (ValueError, IndexError):
                pass
        return None


    @staticmethod
    def is_list_item(paragraph: Paragraph, max_depth: int = 12):

        if not isinstance(paragraph, Paragraph):
            return False, None

        p = paragraph._p
        pPr = p.pPr

        if pPr is not None and pPr.numPr is not None:
            ilvl_nodes = p.xpath('.//w:ilvl/@w:val')
            if ilvl_nodes:
                try:
                    level = int(ilvl_nodes[0]) + 1
                    return True, level
                except (ValueError, IndexError):
                    pass
            return True, 1

        style = paragraph.style
        depth = 0

        while style is not None and depth < max_depth:
            style_elm = style._element
            if style_elm is not None:
                pPr_style = style_elm.pPr
                if pPr_style is not None and pPr_style.numPr is not None:
                    # Нашли numPr в стиле
                    ilvl_nodes_style = style_elm.xpath('.//w:ilvl/@w:val')
                    if ilvl_nodes_style:
                        try:
                            level = int(ilvl_nodes_style[0]) + 1
                            return True, level
                        except (ValueError, IndexError):
                            pass
                    return True, 1  # ilvl отсутствует → уровень 1

            style = style.base_style
            depth += 1

        return False, 0


    @staticmethod
    def has_image(paragraph):
        for run in paragraph.runs:
            xml = str(run.element.xml).lower()
            if 'pic:pic' in xml:
                return True
        return False
    
    @staticmethod
    def has_image_run(run):
        """Проверяет, содержит ли именно этот run изображение"""
        if not run:
            return False
        xml = str(run.element.xml).lower()
        return 'pic:pic' in xml

    def collect_paragraph_errors(self, paragraph, checklist, document):
        """Проверяет каждый run в параграфе и собирает ВСЕ ошибки.
        Возвращает список ошибок для одного общего комментария (дубликаты по параметру удаляются)."""
        if not isinstance(paragraph, Paragraph) or not paragraph.runs:
            return []

        all_errors = []
        seen = set()
        
        paragraph_stats = self.get_paragraph_properties(document, paragraph)

        for run in paragraph.runs:
            if not run.text or not run.text.strip():
                continue

            stats = self.get_run_properties(document, paragraph, run)
            stats.update(paragraph_stats)
            errors = self.get_error_comment(checklist, stats)
            
            for err in errors:
                topic = err.topic
                if topic not in seen:
                    seen.add(topic)
                    all_errors.append(err)

        return all_errors


    def classify_paragraph(self, paragraph):
        """Базовая классификация без учёта контекста (контекст обрабатывается позже)"""
        if not isinstance(paragraph, Paragraph):
            return "unknown", None

        heading_level = self.get_heading_level(paragraph)
        is_list_flag, list_level = self.is_list_item(paragraph)
        has_img = self.has_image(paragraph)

        style_name = (paragraph.style.name or "").lower()

        if heading_level == 1:
            return "heading1", None
        if heading_level == 2:
            return "heading2", None
        if heading_level == 3:
            return "heading3", None
        if heading_level is not None:
            return f"heading{heading_level}", None

        if is_list_flag:
            return "list", list_level

        if has_img:
            return "image", None

        if "caption" in style_name or "подпись" in style_name or "figure" in style_name or "table" in style_name:
            return "caption", None

        return "text", None

    
    def parse_document(self, filename):
        document = Document(filename)
        written_comments = []
        comment_count = 0

        # Состояния
        image_name_check = 0
        text_after_table_check = 0
        first_paragraph_not_reached = True

        # Контекст предыдущего блока
        prev_block_type = None
        prev_paragraph = None

        current_paragraph = None
        current_errors = []
        current_author = "System"
        current_target_runs = []

        for block in self.iter_block_items(document):
            block_errors = []
            block_author = "System"
            target_paragraph = None
            target_runs = []

            current_block_type = "unknown"

            # Поля документа (один раз)
            if first_paragraph_not_reached and isinstance(block, Paragraph):
                first_paragraph_not_reached = False
                margin_comments = self.parse_margins(document)
                for section_title, errors in margin_comments.items():
                    if errors:
                        self.create_comment(document, block.runs, section_title, errors)
                        comment_count += 1
                        written_comments.append([section_title, errors])

            # Параграф
            if isinstance(block, Paragraph):
                p = block
                target_paragraph = p
                base_type, extra = self.classify_paragraph(p)

                current_block_type = base_type

                checklist = None

                # Проверяем, не является ли предыдущий параграф "абзацем перед списком"
                if base_type == "list" and prev_block_type == "text" and prev_paragraph:
                    current_author = "Абзац перед списком"
                    current_errors = self.collect_paragraph_errors(current_paragraph, self.text_before_list_checklist, document)

                # Обычная классификация текущего параграфа
                if base_type.startswith("heading"):
                    level = int(base_type.replace("heading", ""))
                    checklist = getattr(self, f"heading{level}_checklist", self.heading3_checklist)
                    block_author = f"Заголовок {level}"

                elif base_type == "list":
                    paragraph_stats = self.get_paragraph_properties(document, p)
                    if paragraph_stats['is_list']:
                        if paragraph_stats['list_type'] == 'bulleted':
                            checklist = self.bullet_list_checklist
                            block_author = "Маркированный список"
                        elif paragraph_stats['list_type'] == 'numbered':
                            checklist = self.bullet_list_checklist
                            block_author = "Нумерованный список"
                        else:
                            checklist = self.bullet_list_checklist
                            block_author = "Ошибка: список не найден"

                elif base_type == "image":
                    checklist = self.image_checklist
                    block_author = "Рисунок"
                    image_name_check = 2
                    image_run = next((r for r in p.runs if self.has_image_run(r)), None)
                    target_runs = [image_run] if image_run else p.runs

                elif (image_name_check > 0 or base_type == "caption") and prev_block_type == "image":
                    checklist = self.image_name_checklist
                    block_author = "Подпись рисунка"
                    stats = {"format_regex": (p.text or "").strip()}
                    block_errors.extend(self.get_error_comment(checklist, stats))
                    target_runs = p.runs

                elif self.enable_optional_settings.get("paragraph_after_table", False) and text_after_table_check > 0:
                    checklist = self.text_after_table_checklist
                    block_author = "Абзац после таблицы"

                else:
                    checklist = self.text_checklist
                    block_author = "Абзац"

                # Проверка текущего параграфа
                if checklist and 'checklist' in locals() and not block_errors:
                    block_errors = self.collect_paragraph_errors(p, checklist, document)
                    

                if not target_runs:
                    target_runs = p.runs

            # Таблица
            elif isinstance(block, Table):
                
                
                if prev_block_type == "text" and prev_paragraph:
                    current_author = "Название таблицы"
                    current_errors = self.collect_paragraph_errors(current_paragraph, self.table_name_checklist, document)
                    
                current_block_type = "table"
                text_after_table_check = 2
                block_author = "Таблица"

                # Содержимое ячеек (без изменений)
                for row_idx, row in enumerate(block.rows):
                    for col_idx, cell in enumerate(row.cells):
                        if not cell.text.strip():
                            continue
                        for cell_p in cell.paragraphs:
                            if target_paragraph is None:
                                target_paragraph = cell_p
                                target_runs = cell_p.runs if cell_p.runs else []
                                
                            cell_paragraph_stats = self.get_paragraph_properties(document, cell_p)
                            cell_stats = self.get_run_properties(document, cell_p,
                                                                cell_p.runs[0] if cell_p.runs else None)
                            if not cell_stats:
                                continue
                            cell_stats.update(cell_paragraph_stats)
                            cell_stats["vert_alignment"] = cell.vertical_alignment or WD_ALIGN_VERTICAL.TOP

                            checklist_cell = self.table_headings_checklist if \
                                (self.enable_optional_settings.get("table_headings_top", True) and row_idx == 0) or \
                                (self.enable_optional_settings.get("table_headings_left", False) and col_idx == 0) \
                                else self.table_text_checklist

                            cell_errors = self.get_error_comment(checklist_cell, cell_stats)
                            if cell_errors:
                                block_errors.extend(cell_errors)

            # Сохранение предыдущего комментария
            if current_paragraph and current_errors:
                runs_to_use = current_target_runs if current_target_runs else current_paragraph.runs
                current_errors_filtered = []
                seen = set()
                for err in current_errors:
                    topic = err.topic
                    if topic not in seen:
                        seen.add(topic)
                        current_errors_filtered.append(err)
                self.create_comment(document, runs_to_use, current_author, current_errors_filtered)
                comment_count += 1
                written_comments.append([current_author, current_errors_filtered])

            # Обновляем контекст
            prev_block_type = current_block_type
            prev_paragraph = target_paragraph if isinstance(block, Paragraph) else prev_paragraph

            current_paragraph = target_paragraph
            current_errors = block_errors
            current_author = block_author
            current_target_runs = target_runs

            image_name_check = max(0, image_name_check - 1)
            text_after_table_check = max(0, text_after_table_check - 1)

        # Последний блок
        if current_paragraph and current_errors:
            runs_to_use = current_target_runs if current_target_runs else current_paragraph.runs
            current_errors_filtered = []
            seen = set()
            for err in current_errors:
                topic = err.topic
                if topic not in seen:
                    seen.add(topic)
                    current_errors_filtered.append(err)
            self.create_comment(document, runs_to_use, current_author, current_errors_filtered)
            comment_count += 1
            written_comments.append([current_author, current_errors_filtered])

        # Сохранение
        basename = os.path.splitext(os.path.basename(filename))[0]
        out_path = f"Results/{basename}_Проверенный.docx"
        count = 1
        while os.path.exists(out_path):
            out_path = f"Results/{basename}_Проверенный_{count}.docx"
            count += 1

        os.makedirs("Results", exist_ok=True)
        document.save(out_path)

        return comment_count, out_path
    
    
if __name__ == '__main__':
    parser = DocumentParser()
    '''parser.set_settings(default_text_checklist, default_heading1_checklist, default_heading2_checklist,
                        default_heading3_checklist, default_table_name_checklist, default_table_headings_checklist,
                        default_table_text_checklist, default_list_checklist, default_margins_checklist,
                        default_image_checklist, default_image_name_checklist)'''

    filename = "testDocs/testTables.docx"
    parser.parse_document(filename)
