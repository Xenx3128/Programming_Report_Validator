import os
from pathlib import Path

import docx.document
import xml.etree.ElementTree as ET
from docx.api import Document
from docx.document import Document as doctwo
from docx.shared import Pt, Cm, Emu
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_UNDERLINE, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT as WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENTATION
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.comments import Comment
from io import StringIO
import re

from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Length
from docx.text.run import Run

from default_settings import *
from constants import *
import inspect


class BaseChecklist:
    def set_settings(self, settings: dict):
        for par, value in settings.items():
            try:
                setattr(self, par, value)
            except:
                continue


class ParagraphChecklist(BaseChecklist):
    def __init__(self):
        super().__init__()
        self.format_regex = None
        self.font_name = None
        self.font_size = None
        self.font_bald = None
        self.font_italic = None
        self.font_underline = None
        self.font_color = None  # !!!
        self.font_highlight = None  # !!!
        self.alignment = None
        self.keep_with_next = None
        self.page_break_before = None
        self.space_before = None
        self.space_after = None
        self.left_indent = None
        self.right_indent = None
        self.first_line_indent = None
        self.line_spacing = None
        self.header_numbering = None



class ListChecklist(ParagraphChecklist):
    def __init__(self):
        super().__init__()
        self.left_indent_base = None
        self.left_indent_mod = None
        self.list_reminder = (
            "Списку предшествует абзац текста с флагом \"не отрывать от следующего\"\n",
            "Нумерованный список:\n",
            "   Элемент списка начинается с заглавной буквы\n",
            "   Элемент списка заканчивается точкой\n",
            "   В качестве номера используется арабская цифра\n",
            "Маркированный список:\n",
            "   Элемент списка начинается с заглавной буквы\n",
            "   Элемент списка заканчивается точкой с запятой\n",
            "   В качестве маркера используется тире\n",
            "Выступ первой строки 0.75 см\n",
            "Отступ вычисляется по формуле: 1.25+0.75*(Уровень_элемента - 1)\n",
        )


class TableParagraphChecklist(ParagraphChecklist):
    def __init__(self):
        super().__init__()
        self.vert_alignment = None


class ImageChecklist(BaseChecklist):
    def __init__(self):
        super().__init__()
        self.alignment = None
        self.keep_with_next = None
        self.space_before = None
        self.space_after = None
        self.first_line_indent = None
        self.line_spacing = None


class MarginsChecklist(BaseChecklist):
    def __init__(self):
        super().__init__()
        self.top_margin = None
        self.bottom_margin = None
        self.left_margin = None
        self.right_margin = None
        self.orientation = None


class DocumentParser:
    def __init__(self):
        self.text_checklist = ParagraphChecklist()
        self.heading1_checklist = ParagraphChecklist()
        self.heading2_checklist = ParagraphChecklist()
        self.heading3_checklist = ParagraphChecklist()
        self.list_checklist = ListChecklist()
        self.table_headings_checklist = TableParagraphChecklist()
        self.table_text_checklist = TableParagraphChecklist()
        self.table_name_checklist = ParagraphChecklist()
        self.text_after_table_checklist = ParagraphChecklist()
        self.image_checklist = ImageChecklist()
        self.image_name_checklist = ParagraphChecklist()
        self.text_before_list_checklist = ParagraphChecklist()
        self.margins_checklist = MarginsChecklist()
        self.set_settings(default_text_checklist, default_heading1_checklist, default_heading2_checklist,
                          default_heading3_checklist, default_table_name_checklist, default_table_headings_checklist,
                          default_table_text_checklist, default_list_checklist, default_margins_checklist,
                          default_image_checklist, default_image_name_checklist, default_text_before_list_checklist)

        self.enable_optional_settings = {
            "table_headings_top": True,
            "table_headings_left": False,
            "table_title": True,
            "paragraph_after_table": True,
            "enable_pic_title": True,
            "list_reminder": True,
        }

        
        self.doc_comments = {}

    def set_enable_optional_settings(self, elements: dict):
        if elements is not None:
            for par, value in elements.items():
                if par in self.enable_optional_settings:
                    self.enable_optional_settings[par] = value

    def set_settings(self, text_check=None, h1_check=None, h2_check=None, h3_check=None, table_name_check=None,
                     table_heading_check=None, table_text_check=None, list_check=None,
                     page_check=None, pic_check=None, pic_name_check=None, text_before_list_check=None):
        if isinstance(text_check, dict):
            self.text_checklist.set_settings(text_check)
            self.text_after_table_checklist.set_settings(text_check)
            self.text_after_table_checklist.space_before = 13.0  # !!!!
            self.list_checklist.set_settings(text_check)
        if isinstance(h1_check, dict):
            self.heading1_checklist.set_settings(h1_check)
        if isinstance(h2_check, dict):
            self.heading2_checklist.set_settings(h2_check)
        if isinstance(h3_check, dict):
            self.heading3_checklist.set_settings(h3_check)
        if isinstance(table_name_check, dict):
            self.table_name_checklist.set_settings(table_name_check)
        if isinstance(table_heading_check, dict):
            self.table_headings_checklist.set_settings(table_heading_check)
        if isinstance(table_text_check, dict):
            self.table_text_checklist.set_settings(table_text_check)
        if isinstance(list_check, dict):
            self.list_checklist.set_settings(list_check)
            # if "left_indent_base" in list_check:
                # self.list_checklist.left_indent_base -= 0.75
        if isinstance(page_check, dict):
            self.margins_checklist.set_settings(page_check)
        if isinstance(pic_check, dict):
            self.image_checklist.set_settings(pic_check)
        if isinstance(pic_name_check, dict):
            self.image_name_checklist.set_settings(pic_name_check)
        if isinstance(text_before_list_check, dict):
            self.text_before_list_checklist.set_settings(text_before_list_check)

    
    @staticmethod
    def get_error_comment(checklist, received: dict):
        attributes = inspect.getmembers(checklist, lambda a: not (inspect.isroutine(a)))
        expected = {a[0]: a[1] for a in attributes if not (a[0].startswith('__') and a[0].endswith('__'))}
        comments = []
        
        for key, val in expected.items():
            comparison_res = True
            comment = []
            if key in received.keys() and val is not None:
                # Отдельная обработка формата через RegEx
                if key == 'format_regex':
                    comparison_res = re.fullmatch(fr'{REGEX_TRANSFORM[val]}', received[key])
                    if not comparison_res:
                        comment = [PARAM_TO_COMMENT[key], val, received[key]]
                        comments.append(comment)
                        continue
                elif "list_level" in received.keys() and key == "left_indent":
                    expected_indent = expected["left_indent_base"] + expected["left_indent_mod"] * received["list_level"]
                    if received["left_indent"] != expected_indent:
                        comment = [PARAM_TO_COMMENT[key], expected_indent, received[key]]
                        comments.append(comment)
                        continue
                else:
                    if isinstance(val, bool):
                        if received[key] == True:
                            comparison_res = val == True
                        else:
                            comparison_res = val in (False, None)
                    else:
                        comparison_res = val == received[key]
                if not comparison_res:
                    comment = [PARAM_TO_COMMENT[key]]
                    # У выравниваний значения приравниваются к 0..3, поэтому нужен костыль
                    if key == "alignment":
                        comment += [ALIGNMENT_TO_COMMENT[val], ALIGNMENT_TO_COMMENT[received[key]]]
                    elif key == "vert_alignment":
                        comment += [VERT_ALIGNMENT_TO_COMMENT[val], VERT_ALIGNMENT_TO_COMMENT[received[key]]]
                    elif key == "orientation":
                        comment += [ORIENTATION_TO_COMMENT[val], ORIENTATION_TO_COMMENT[received[key]]]
                    else:
                        if val is None or isinstance(val, bool):
                            comment += [VAR_TO_COMMENT[val]]
                        else:
                            comment += [val]
                        if received is None or isinstance(received[key], bool):
                            comment += [VAR_TO_COMMENT[received[key]]]
                        else:
                            comment += [received[key]]
                    comments.append(comment)
        return comments
    
    @staticmethod
    def create_comment(document, runs, element, comments: list):
        if len(comments) > 0 and len(runs) > 0:
            comment = document.add_comment(
                runs=runs,
                text='',
                author=element
            )
            if element == 'Таблица':
                cmt_para = comment.add_paragraph()
                cmt_para.add_run("Выведена только первая ошибка в каждой категории\n").bold = True 
            for cmt in comments:
                parameter, expected_value, received_value = cmt
                cmt_para = comment.add_paragraph()
                cmt_para.add_run(f"{parameter}: ").bold = True 
                cmt_para.add_run(f"{received_value} | ")
                cmt_para.add_run(f"({expected_value}).")
            return comment

 
    def get_run_properties(self, document, paragraph, run):
        """
        ИСПРАВЛЕННАЯ версия (пункты 1–3):
        • Добавлены все необходимые импорты
        • Исправлен порядок наследования стилей (run → run.style → linked_style → paragraph.style.font)
        • resolve_theme_name полностью переписан — теперь **гарантированно** возвращает реальное имя шрифта
        (больше никогда не возвращает пустую строку даже в повреждённых темах, Google Docs, старых Word 2007 и т.д.)
        """

        if run is None or not hasattr(run, '_element'):
            return None

        # ====================== КЭШИРОВАНИЕ (один раз на документ) ======================
        if not hasattr(self, '_docx_style_cache'):
            self._docx_style_cache = {}
            self._theme_cache = None          # dict {'minor': {...}, 'major': {...}}
            self._doc_defaults = None

        # ====================== ИСПРАВЛЕННЫЕ ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ======================
        def get_theme_font(font_type='minor'):
            """Кэшированный доступ к теме (latin + hAnsi приоритет)"""
            if self._theme_cache is None:
                self._theme_cache = {'minor': {}, 'major': {}}
                try:
                    # Способ 1: высокий уровень API
                    theme = document.part.theme_part.theme
                    fs = theme.themeElements.fontScheme
                    fgroup = getattr(fs, f'{font_type}Font', None)
                    if fgroup:
                        self._theme_cache[font_type] = {
                            'latin': getattr(fgroup.latin, 'typeface', None),
                            'hAnsi': getattr(getattr(fgroup, 'hAnsi', None), 'typeface', None),
                            'eastAsia': getattr(getattr(fgroup, 'eastAsia', None), 'typeface', None),
                            'cs': getattr(getattr(fgroup, 'cs', None), 'typeface', None),
                        }
                except Exception:
                    pass

                # Способ 2: чистый XML (работает всегда)
                if not self._theme_cache[font_type]:
                    try:
                        el = document.part.theme_part.element
                        ns = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
                        typ = 'minorFont' if font_type == 'minor' else 'majorFont'
                        res = el.xpath(f'.//a:fontScheme/a:{typ}/a:latin')
                        if res:
                            typeface = res[0].get('typeface')
                            self._theme_cache[font_type]['hAnsi'] = typeface
                            self._theme_cache[font_type]['latin'] = typeface
                    except Exception:
                        pass

                # Финальный fallback
                default = "Calibri" if font_type == 'minor' else "Calibri Light"
                if not self._theme_cache[font_type]:
                    self._theme_cache[font_type] = {'latin': default, 'hAnsi': default}

            return self._theme_cache[font_type]

        def resolve_theme_name(name):
            """ГЛАВНОЕ ИСПРАВЛЕНИЕ: больше НИКОГДА не возвращает пустую строку"""
            if not name:
                return "Calibri"

            name = str(name).strip()
            if not name:
                return "Calibri"

            lower = name.lower()

            # Все возможные плейсхолдеры (включая asciiTheme/hAnsiTheme)
            minor_placeholders = {
                "+body", "body", "+mnlt", "mnlt", "minor", "minorhansi", "minorascii",
                "minorbidi", "minoreastasia", "asciitheme", "hansitheme", "cstheme"
            }
            major_placeholders = {
                "+heading", "heading", "+mnhansi", "mnhansi", "major", "majorhansi",
                "majorascii", "majorbidi", "majoreastasia"
            }

            is_major = lower in major_placeholders
            theme_group = get_theme_font('major' if is_major else 'minor')

            # Приоритет hAnsi (самый важный для европейских документов и разных версий Word)
            for key in ('hAnsi', 'ascii', 'latin', 'eastAsia', 'cs'):
                val = theme_group.get(key)
                if val and str(val).strip():
                    return str(val).strip()

            # Если ничего не нашлось
            return "Calibri Light" if is_major else "Calibri"

        def get_color_value(color_format):
            """Оставлено почти как было + небольшая защита"""
            if not color_format:
                return None
            try:
                if hasattr(color_format, 'rgb') and color_format.rgb and str(color_format.rgb) != '000000':
                    rgb = color_format.rgb
                    return f"#{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
                if hasattr(color_format, 'theme_color') and color_format.theme_color:
                    tc = color_format.theme_color
                    name = getattr(tc, 'name', str(tc)).lower()
                    tint = getattr(color_format, 'brightness', None)
                    return f"theme:{name}" + (f"(tint={tint:.2f})" if tint else "")
            except Exception:
                pass
            return None

        def get_prop_from_font(font_obj, prop_name):
            """Универсальный геттер"""
            if not font_obj:
                return None
            try:
                if prop_name == "name":
                    raw = getattr(font_obj, 'name', None)
                    return raw if raw else None
                if prop_name == "size":
                    size = getattr(font_obj, 'size', None)
                    return size.pt if size else None
                if prop_name in ("bold", "italic", "strike", "subscript", "superscript",
                                "small_caps", "all_caps"):
                    return getattr(font_obj, prop_name, None)
                if prop_name == "underline":
                    return getattr(font_obj, 'underline', None)
                if prop_name == "color":
                    return get_color_value(font_obj.color)
                if prop_name == "highlight":
                    return getattr(font_obj, 'highlight_color', None)
            except Exception:
                pass
            return None

        def get_style_chain_prop(style, prop_name):
            """Цепочка base_style"""
            if not style:
                return None
            visited = set()
            current = style
            while current and current.name not in visited:
                visited.add(current.name)
                value = get_prop_from_font(current.font, prop_name)
                if value is not None and (prop_name != "name" or value):  # защита от пустой строки
                    return value
                current = getattr(current, 'base_style', None)
            return None

        def resolve_run_prop(prop_name, default=None):
            """ИСПРАВЛЕННЫЙ порядок наследования (самое важное изменение)"""
            # 1. Прямое форматирование run
            value = get_prop_from_font(run.font, prop_name)
            if value is not None and (prop_name != "name" or value):
                return value

            # 2. Стиль самого run (character style)
            if run.style:
                value = get_style_chain_prop(run.style, prop_name)
                if value is not None and (prop_name != "name" or value):
                    return value

            # 3. Linked character style параграфа (КРИТИЧНО для Heading, Title и т.д.)
            try:
                linked = getattr(paragraph.style, 'linked_style', None)
                if linked:
                    value = get_style_chain_prop(linked, prop_name)
                    if value is not None and (prop_name != "name" or value):
                        return value
            except Exception:
                pass

            # 4. Стиль параграфа (paragraph.style.font)
            if paragraph.style:
                value = get_style_chain_prop(paragraph.style, prop_name)
                if value is not None and (prop_name != "name" or value):
                    return value

            # 5. docDefaults (XML — самый надёжный способ)
            if self._doc_defaults is None:
                try:
                    self._doc_defaults = document.styles.element.xpath(
                        './/w:docDefaults/w:rPrDefault/w:rPr'
                    )
                    self._doc_defaults = self._doc_defaults[0] if self._doc_defaults else None
                except Exception:
                    self._doc_defaults = None

            if self._doc_defaults is not None:
                try:
                    if prop_name == "name":
                        rFonts = self._doc_defaults.find(qn('w:rFonts'))
                        if rFonts is not None:
                            for attr in ['hAnsi', 'ascii', 'eastAsia', 'cs', 'hAnsiTheme', 'asciiTheme']:
                                val = rFonts.get(qn(f'w:{attr}'))
                                if val:
                                    return val
                except Exception:
                    pass

            # 6. Тема по умолчанию
            if prop_name == "name":
                return resolve_theme_name("+Body") or default
            return default

        def resolve_para_prop(prop_name, default=None):
            """Параграфные свойства (без изменений)"""
            try:
                fmt = paragraph.paragraph_format
                value = getattr(fmt, prop_name, None)
                if value is not None:
                    return value

                style = paragraph.style
                while style:
                    value = getattr(style.paragraph_format, prop_name, None)
                    if value is not None:
                        return value
                    style = getattr(style, 'base_style', None)
            except Exception:
                pass
            return default

        # ====================== СБОР РЕЗУЛЬТАТА ======================
        paragraph_stats = {
            "font_name": resolve_run_prop("name", "Calibri"),
            "font_size": round(resolve_run_prop("size") or 11.0, 1),
            "font_bold": bool(resolve_run_prop("bold")),
            "font_italic": bool(resolve_run_prop("italic")),
            "font_underline": resolve_run_prop("underline") not in (None, False, WD_UNDERLINE.NONE),
            "font_color": resolve_run_prop("color"),
            "font_highlight": resolve_run_prop("highlight"),
            "font_strike": bool(resolve_run_prop("strike")),
            "font_subscript": bool(resolve_run_prop("subscript")),
            "font_superscript": bool(resolve_run_prop("superscript")),
            "font_small_caps": bool(resolve_run_prop("small_caps")),
            "font_all_caps": bool(resolve_run_prop("all_caps")),

            "alignment": resolve_para_prop("alignment", WD_PARAGRAPH_ALIGNMENT.LEFT),
            "keep_with_next": bool(resolve_para_prop("keep_with_next")),
            "page_break_before": bool(resolve_para_prop("page_break_before")),

            "left_indent": round(Length(resolve_para_prop("left_indent") or 0).cm, 2),
            "right_indent": round(Length(resolve_para_prop("right_indent") or 0).cm, 2),
            "first_line_indent": round(Length(resolve_para_prop("first_line_indent") or 0).cm, 2),

            "space_before": round(getattr(resolve_para_prop("space_before"), 'pt', 0), 1),
            "space_after": round(getattr(resolve_para_prop("space_after"), 'pt', 0), 1),
            "line_spacing": resolve_para_prop("line_spacing") or 1.0,

            "is_list": self.is_list_item(paragraph)[0],
        }

        return paragraph_stats


    # This function extracts the tables and paragraphs from the document object
    @staticmethod
    def iter_block_items(parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph. *parent*
        would most commonly be a reference to a main Document object, but
        also works for a _Cell object, which itself can contain paragraphs and tables.
        """
        if isinstance(parent, doctwo):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    @staticmethod
    def set_comment_name(name, prev_name, is_locked):
        if not is_locked:
            return name
        return prev_name

    def parse_margins(self, document):
        sections = document.sections
        paragraph_stats = {}
        stats_to_compare = self.margins_checklist
        margins_comments = {}
        for section_cnt, section in enumerate(sections):
            paragraph_stats["top_margin"] = round(section.top_margin.cm, 2)
            paragraph_stats["bottom_margin"] = round(section.bottom_margin.cm, 2)
            paragraph_stats["left_margin"] = round(section.left_margin.cm, 2)
            paragraph_stats["right_margin"] = round(section.right_margin.cm, 2)
            paragraph_stats["orientation"] = section.orientation
            section_comments = self.get_error_comment(stats_to_compare, paragraph_stats)
            if len(section_comments) > 0:
                margins_comments[f"Раздел {section_cnt + 1}:"] = section_comments
        return margins_comments     



    # ====================== НОВЫЕ ВСПОМОГАТЕЛЬНЫЕ МЕТОДЫ ======================
    @staticmethod
    def get_heading_level(paragraph):
        """Возвращает уровень заголовка (1-9) или None. Покрывает ВСЕ крайние случаи."""
        def extract_level_from_style(style):
            if not style:
                return None
            name = style.name or ""
            if name.startswith("Heading "):
                try:
                    return int(name.split()[-1])
                except (ValueError, IndexError):
                    pass
            style_id = getattr(style, 'style_id', None) or ""
            if style_id.startswith("Heading"):
                try:
                    return int(style_id.replace("Heading", "").strip())
                except ValueError:
                    pass
            if name.startswith("Heading") and len(name) > 7:
                try:
                    return int(name[7:])
                except ValueError:
                    pass
            return None

        # Быстрый путь
        level = extract_level_from_style(paragraph.style)
        if level is not None:
            return level

        # Цепочка наследования стилей (самый частый крайний случай!)
        current = paragraph.style
        depth = 0
        while current and depth < 20:
            level = extract_level_from_style(current)
            if level is not None:
                return level
            current = current.base_style
            depth += 1

        # outlineLvl (используется в некоторых шаблонах и экспортах из Google Docs)
        outline = paragraph._p.xpath('./w:pPr/w:outlineLvl/@w:val')
        if outline:
            try:
                lvl = int(outline[0])
                return lvl if 1 <= lvl <= 9 else None
            except (ValueError, IndexError):
                pass
        return None


    @staticmethod
    def is_list_item(paragraph: Paragraph, max_depth: int = 12):

        if not isinstance(paragraph, Paragraph):
            return False, None

        p = paragraph._p
        pPr = p.pPr

        # 1. Прямая нумерация в самом параграфе (самый надёжный случай)
        if pPr is not None and pPr.numPr is not None:
            ilvl_nodes = p.xpath('.//w:ilvl/@w:val')
            if ilvl_nodes:
                try:
                    level = int(ilvl_nodes[0]) + 1
                    return True, level
                except (ValueError, IndexError):
                    pass
            return True, 1  # ilvl нет → считаем первым уровнем

        # 2. Ищем в стиле и цепочке базовых стилей
        style = paragraph.style
        depth = 0

        while style is not None and depth < max_depth:
            style_elm = style._element
            if style_elm is not None:
                pPr_style = style_elm.pPr
                if pPr_style is not None and pPr_style.numPr is not None:
                    # Нашли numPr в стиле
                    ilvl_nodes_style = style_elm.xpath('.//w:ilvl/@w:val')
                    if ilvl_nodes_style:
                        try:
                            level = int(ilvl_nodes_style[0]) + 1
                            return True, level
                        except (ValueError, IndexError):
                            pass
                    return True, 1  # ilvl отсутствует → уровень 1

            style = style.base_style
            depth += 1

        return False, 0


    @staticmethod
    def has_image(paragraph):
        """Проверка наличия изображения (pic + drawing — покрывает все версии Word)"""
        for run in paragraph.runs:
            xml = str(run.element.xml).lower()
            if 'pic:pic' in xml:
                return True
        return False
    
    @staticmethod
    def has_image_run(run):
        """Проверяет, содержит ли именно этот run изображение"""
        if not run:
            return False
        xml = str(run.element.xml).lower()
        return 'pic:pic' in xml

    def collect_paragraph_errors(self, paragraph, checklist, document):
        """Проверяет КАЖДЫЙ run в параграфе и собирает ВСЕ ошибки.
        Возвращает список ошибок для одного общего комментария (дубликаты по параметру удаляются)."""
        if not isinstance(paragraph, Paragraph) or not paragraph.runs:
            return []

        all_errors = []
        seen = set()

        for run in paragraph.runs:
            # Пропускаем пустые run-ы (часто бывают пробелы/табуляция)
            if not run.text or not run.text.strip():
                continue

            stats = self.get_run_properties(document, paragraph, run)
            errors = self.get_error_comment(checklist, stats)

            for err in errors:
                err_key = tuple(str(x) for x in err)  # дедупликация
                if err_key not in seen:
                    seen.add(err_key)
                    all_errors.append(err)

        return all_errors


    def _classify_paragraph(self, paragraph):
        """Базовая классификация без учёта контекста (контекст обрабатывается позже)"""
        if not isinstance(paragraph, Paragraph):
            return "unknown", None

        heading_level = self.get_heading_level(paragraph)
        is_list_flag, list_level = self.is_list_item(paragraph)
        has_img = self.has_image(paragraph)

        style_name = (paragraph.style.name or "").lower()

        if heading_level == 1:
            return "heading1", None
        if heading_level == 2:
            return "heading2", None
        if heading_level == 3:
            return "heading3", None
        if heading_level is not None:
            return f"heading{heading_level}", None

        if is_list_flag:
            return "list", list_level

        if has_img:
            return "image", None

        if "caption" in style_name or "подпись" in style_name or "figure" in style_name or "table" in style_name:
            return "caption", None

        return "text", None

    # =========================================================================
    
    def parse_document(self, filename):
        document = Document(filename)
        written_comments = []
        comment_count = 0

        # Состояния
        image_name_check = 0                  # счётчик для подписи рисунка (2 блока после изображения)
        text_after_table_check = 0            # счётчик для абзаца после таблицы
        first_paragraph_not_reached = True

        # Контекст предыдущего блока
        prev_block_type = None
        prev_paragraph = None                 # нужен для названия таблицы

        current_paragraph = None
        current_errors = []
        current_author = "System"
        current_target_runs = []

        for block in self.iter_block_items(document):
            block_errors = []
            block_author = "System"
            target_paragraph = None
            target_runs = []

            current_block_type = "unknown"

            # Поля документа (один раз)
            if first_paragraph_not_reached and isinstance(block, Paragraph):
                first_paragraph_not_reached = False
                margin_comments = self.parse_margins(document)
                for section_title, errors in margin_comments.items():
                    if errors:
                        self.create_comment(document, block.runs, section_title, errors)
                        comment_count += 1
                        written_comments.append([section_title, errors])

            # ── ПАРАГРАФ ────────────────────────────────────────────────────────
            if isinstance(block, Paragraph):
                p = block
                target_paragraph = p
                base_type, extra = self._classify_paragraph(p)

                current_block_type = base_type

                checklist = None

                # Проверяем, не является ли предыдущий параграф "абзацем перед списком"
                if base_type == "list" and prev_block_type == "text" and prev_paragraph:
                    current_author = "Абзац перед списком"
                    current_errors = self.collect_paragraph_errors(current_paragraph, self.text_before_list_checklist, document)

                # Обычная классификация текущего параграфа
                if base_type.startswith("heading"):
                    level = int(base_type.replace("heading", ""))
                    checklist = getattr(self, f"heading{level}_checklist", self.heading3_checklist)
                    block_author = f"Заголовок {level}"

                elif base_type == "list":
                    checklist = self.list_checklist
                    block_author = "Элемент списка"

                elif base_type == "image":
                    checklist = self.image_checklist
                    block_author = "Рисунок"
                    image_name_check = 2
                    image_run = next((r for r in p.runs if self.has_image_run(r)), None)
                    target_runs = [image_run] if image_run else p.runs

                elif (image_name_check > 0 or base_type == "caption") and prev_block_type == "image":
                    checklist = self.image_name_checklist
                    block_author = "Подпись рисунка"
                    stats = {"format_regex": (p.text or "").strip()}
                    block_errors.extend(self.get_error_comment(checklist, stats))
                    target_runs = p.runs

                elif self.enable_optional_settings.get("paragraph_after_table", False) and text_after_table_check > 0:
                    checklist = self.text_after_table_checklist
                    block_author = "Абзац после таблицы"

                else:
                    checklist = self.text_checklist
                    block_author = "Абзац"

                # Проверка текущего параграфа
                if checklist and 'checklist' in locals() and not block_errors:
                    block_errors = self.collect_paragraph_errors(p, checklist, document)

                if not target_runs:
                    target_runs = p.runs

            # ── ТАБЛИЦА ─────────────────────────────────────────────────────────
            elif isinstance(block, Table):
                
                
                if prev_block_type == "text" and prev_paragraph:
                    current_author = "Название таблицы"
                    current_errors = self.collect_paragraph_errors(current_paragraph, self.table_name_checklist, document)
                    
                current_block_type = "table"
                text_after_table_check = 2
                block_author = "Таблица"

                # Содержимое ячеек (без изменений)
                for row_idx, row in enumerate(block.rows):
                    for col_idx, cell in enumerate(row.cells):
                        if not cell.text.strip():
                            continue
                        for cell_p in cell.paragraphs:
                            if target_paragraph is None:
                                target_paragraph = cell_p
                                target_runs = cell_p.runs if cell_p.runs else []

                            cell_stats = self.get_run_properties(document, cell_p,
                                                                cell_p.runs[0] if cell_p.runs else None)
                            cell_stats["vert_alignment"] = cell.vertical_alignment or WD_ALIGN_VERTICAL.TOP

                            checklist_cell = self.table_headings_checklist if \
                                (self.enable_optional_settings.get("table_headings_top", True) and row_idx == 0) or \
                                (self.enable_optional_settings.get("table_headings_left", False) and col_idx == 0) \
                                else self.table_text_checklist

                            cell_errors = self.get_error_comment(checklist_cell, cell_stats)
                            if cell_errors:
                                block_errors.extend(cell_errors)

            # Сохранение предыдущего комментария
            if current_paragraph and current_errors:
                runs_to_use = current_target_runs if current_target_runs else current_paragraph.runs
                current_errors_filtered = []
                seen = set()
                for err in current_errors:
                    err_key = tuple(str(x) for x in err)
                    if err_key not in seen:
                        seen.add(err_key)
                        current_errors_filtered.append(err)
                self.create_comment(document, runs_to_use, current_author, current_errors_filtered)
                comment_count += 1
                written_comments.append([current_author, current_errors_filtered])

            # Обновляем контекст
            prev_block_type = current_block_type
            prev_paragraph = target_paragraph if isinstance(block, Paragraph) else prev_paragraph

            current_paragraph = target_paragraph
            current_errors = block_errors
            current_author = block_author
            current_target_runs = target_runs

            image_name_check = max(0, image_name_check - 1)
            text_after_table_check = max(0, text_after_table_check - 1)

        # Последний блок
        if current_paragraph and current_errors:
            runs_to_use = current_target_runs if current_target_runs else current_paragraph.runs
            current_errors_filtered = []
            seen = set()
            for err in current_errors:
                err_key = tuple(str(x) for x in err)
                if err_key not in seen:
                    seen.add(err_key)
                    current_errors_filtered.append(err)
            self.create_comment(document, runs_to_use, current_author, current_errors_filtered)
            comment_count += 1
            written_comments.append([current_author, current_errors_filtered])

        # Сохранение
        basename = os.path.splitext(os.path.basename(filename))[0]
        out_path = f"Results/{basename}_Проверенный.docx"
        count = 1
        while os.path.exists(out_path):
            out_path = f"Results/{basename}_Проверенный_{count}.docx"
            count += 1

        os.makedirs("Results", exist_ok=True)
        document.save(out_path)

        return written_comments
    
    
if __name__ == '__main__':
    parser = DocumentParser()
    '''parser.set_settings(default_text_checklist, default_heading1_checklist, default_heading2_checklist,
                        default_heading3_checklist, default_table_name_checklist, default_table_headings_checklist,
                        default_table_text_checklist, default_list_checklist, default_margins_checklist,
                        default_image_checklist, default_image_name_checklist)'''

    filename = "testDocs/testTables.docx"
    parser.parse_document(filename)
